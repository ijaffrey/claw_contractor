"""
Contract Extractor Module

This module provides functionality to extract contract data from various sources
including PDF files, Word documents, text files, and database records.
"""

import os
import re
import json
import logging
from datetime import datetime
from typing import Dict, List, Optional, Union, Any
from pathlib import Path

import pandas as pd
from dateutil import parser as date_parser
from PyPDF2 import PdfReader
from docx import Document
import sqlite3
import psycopg2
from psycopg2.extras import RealDictCursor

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class ContractExtractionError(Exception):
    """Custom exception for contract extraction errors"""

    pass


class ContractExtractor:
    """
    Main contract extractor class that handles various data sources
    """

    def __init__(self, config: Optional[Dict] = None):
        """
        Initialize the contract extractor

        Args:
            config: Configuration dictionary with database connections, file paths, etc.
        """
        self.config = config or {}
        self.patterns = self._load_extraction_patterns()

    def _load_extraction_patterns(self) -> Dict[str, re.Pattern]:
        """Load regex patterns for extracting contract information"""
        return {
            "contract_number": re.compile(
                r"(?:contract\s+(?:number|no\.?|#)\s*[:\-]?\s*)([A-Z0-9\-\_]+)",
                re.IGNORECASE,
            ),
            "date": re.compile(
                r"(\d{1,2}[\/\-\.]\d{1,2}[\/\-\.]\d{2,4}|\d{4}[\/\-\.]\d{1,2}[\/\-\.]\d{1,2})",
                re.IGNORECASE,
            ),
            "amount": re.compile(r"(?:\$|USD|EUR|GBP)\s?([\d,]+\.?\d*)", re.IGNORECASE),
            "party_names": re.compile(
                r"(?:party|client|contractor|vendor|supplier)\s*[:\-]?\s*([A-Za-z\s&\.,Inc]+)",
                re.IGNORECASE,
            ),
            "email": re.compile(r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b"),
            "phone": re.compile(r"(?:\+?1[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}"),
            "address": re.compile(
                r"\d+\s+[A-Za-z\s,]+(?:Street|St|Avenue|Ave|Road|Rd|Boulevard|Blvd|Drive|Dr)",
                re.IGNORECASE,
            ),
        }

    def extract_from_pdf(self, file_path: str) -> Dict[str, Any]:
        """
        Extract contract data from PDF files

        Args:
            file_path: Path to the PDF file

        Returns:
            Dictionary containing extracted contract information
        """
        try:
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"PDF file not found: {file_path}")

            with open(file_path, "rb") as file:
                pdf_reader = PdfReader(file)
                text_content = ""

                for page in pdf_reader.pages:
                    text_content += page.extract_text() + "\n"

                if not text_content.strip():
                    logger.warning(f"No text content found in PDF: {file_path}")
                    return {"source": file_path, "error": "No extractable text content"}

                return self._extract_from_text(text_content, source=file_path)

        except Exception as e:
            logger.error(f"Error extracting from PDF {file_path}: {str(e)}")
            raise ContractExtractionError(f"PDF extraction failed: {str(e)}")

    def extract_from_docx(self, file_path: str) -> Dict[str, Any]:
        """
        Extract contract data from Word documents

        Args:
            file_path: Path to the DOCX file

        Returns:
            Dictionary containing extracted contract information
        """
        try:
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"DOCX file not found: {file_path}")

            doc = Document(file_path)
            text_content = ""

            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"

            # Extract from tables if present
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_content += cell.text + " "
                    text_content += "\n"

            if not text_content.strip():
                logger.warning(f"No text content found in DOCX: {file_path}")
                return {"source": file_path, "error": "No extractable text content"}

            return self._extract_from_text(text_content, source=file_path)

        except Exception as e:
            logger.error(f"Error extracting from DOCX {file_path}: {str(e)}")
            raise ContractExtractionError(f"DOCX extraction failed: {str(e)}")

    def extract_from_text_file(
        self, file_path: str, encoding: str = "utf-8"
    ) -> Dict[str, Any]:
        """
        Extract contract data from text files

        Args:
            file_path: Path to the text file
            encoding: File encoding (default: utf-8)

        Returns:
            Dictionary containing extracted contract information
        """
        try:
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"Text file not found: {file_path}")

            with open(file_path, "r", encoding=encoding) as file:
                text_content = file.read()

            if not text_content.strip():
                logger.warning(f"No content found in text file: {file_path}")
                return {"source": file_path, "error": "Empty file"}

            return self._extract_from_text(text_content, source=file_path)

        except Exception as e:
            logger.error(f"Error extracting from text file {file_path}: {str(e)}")
            raise ContractExtractionError(f"Text file extraction failed: {str(e)}")

    def extract_from_database(
        self, query: str, connection_params: Dict[str, str], db_type: str = "postgresql"
    ) -> List[Dict[str, Any]]:
        """
        Extract contract data from database

        Args:
            query: SQL query to retrieve contract data
            connection_params: Database connection parameters
            db_type: Type of database ('postgresql', 'sqlite', 'mysql')

        Returns:
            List of dictionaries containing contract information
        """
        try:
            if db_type.lower() == "postgresql":
                return self._extract_from_postgresql(query, connection_params)
            elif db_type.lower() == "sqlite":
                return self._extract_from_sqlite(query, connection_params)
            else:
                raise ValueError(f"Unsupported database type: {db_type}")

        except Exception as e:
            logger.error(f"Error extracting from database: {str(e)}")
            raise ContractExtractionError(f"Database extraction failed: {str(e)}")

    def extract_from_csv(
        self, file_path: str, mapping: Optional[Dict[str, str]] = None
    ) -> List[Dict[str, Any]]:
        """
        Extract contract data from CSV files

        Args:
            file_path: Path to the CSV file
            mapping: Column name mapping for contract fields

        Returns:
            List of dictionaries containing contract information
        """
        try:
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"CSV file not found: {file_path}")

            df = pd.read_csv(file_path)

            if df.empty:
                logger.warning(f"Empty CSV file: {file_path}")
                return []

            # Apply column mapping if provided
            if mapping:
                df = df.rename(columns=mapping)

            contracts = []
            for index, row in df.iterrows():
                contract_data = {
                    "source": f"{file_path}:row_{index + 1}",
                    "extraction_date": datetime.now().isoformat(),
                    "raw_data": row.to_dict(),
                }

                # Extract standardized fields
                contract_data.update(self._standardize_contract_data(row.to_dict()))
                contracts.append(contract_data)

            return contracts

        except Exception as e:
            logger.error(f"Error extracting from CSV {file_path}: {str(e)}")
            raise ContractExtractionError(f"CSV extraction failed: {str(e)}")

    def _extract_from_text(self, text_content: str, source: str) -> Dict[str, Any]:
        """
        Extract contract information from text content using regex patterns

        Args:
            text_content: Raw text content
            source: Source identifier

        Returns:
            Dictionary containing extracted contract information
        """
        extracted_data = {
            "source": source,
            "extraction_date": datetime.now().isoformat(),
            "raw_text": (
                text_content[:1000] if len(text_content) > 1000 else text_content
            ),  # Truncate for storage
        }

        # Extract contract number
        contract_numbers = self.patterns["contract_number"].findall(text_content)
        if contract_numbers:
            extracted_data["contract_number"] = contract_numbers[0].strip()

        # Extract dates
        date_matches = self.patterns["date"].findall(text_content)
        dates = []
        for date_str in date_matches:
            try:
                parsed_date = date_parser.parse(date_str)
                dates.append(parsed_date.isoformat())
            except:
                continue
        if dates:
            extracted_data["dates"] = dates
            extracted_data["contract_date"] = dates[
                0
            ]  # Assume first date is contract date

        # Extract monetary amounts
        amounts = self.patterns["amount"].findall(text_content)
        if amounts:
            try:
                # Clean and convert to float
                amount_str = amounts[0].replace(",", "")
                extracted_data["contract_amount"] = float(amount_str)
            except:
                extracted_data["contract_amount_raw"] = amounts[0]

        # Extract party names
        parties = self.patterns["party_names"].findall(text_content)
        if parties:
            extracted_data["parties"] = [
                party.strip() for party in parties[:5]
            ]  # Limit to 5 parties

        # Extract contact information
        emails = self.patterns["email"].findall(text_content)
        if emails:
            extracted_data["emails"] = list(set(emails[:5]))  # Unique emails, max 5

        phones = self.patterns["phone"].findall(text_content)
        if phones:
            extracted_data["phone_numbers"] = list(
                set(phones[:3])
            )  # Unique phones, max 3

        addresses = self.patterns["address"].findall(text_content)
        if addresses:
            extracted_data["addresses"] = list(
                set(addresses[:3])
            )  # Unique addresses, max 3

        return extracted_data

    def _extract_from_postgresql(
        self, query: str, connection_params: Dict[str, str]
    ) -> List[Dict[str, Any]]:
        """Extract data from PostgreSQL database"""
        conn = None
        try:
            conn = psycopg2.connect(**connection_params)
            cursor = conn.cursor(cursor_factory=RealDictCursor)
            cursor.execute(query)

            results = []
            for row in cursor.fetchall():
                contract_data = dict(row)
                contract_data["source"] = (
                    f"postgresql:{connection_params.get('host', 'localhost')}"
                )
                contract_data["extraction_date"] = datetime.now().isoformat()
                contract_data.update(self._standardize_contract_data(contract_data))
                results.append(contract_data)

            return results

        finally:
            if conn:
                conn.close()

    def _extract_from_sqlite(
        self, query: str, connection_params: Dict[str, str]
    ) -> List[Dict[str, Any]]:
        """Extract data from SQLite database"""
        db_path = connection_params.get("database", connection_params.get("db_path"))
        if not db_path:
            raise ValueError("SQLite database path not provided")

        conn = None
        try:
            conn = sqlite3.connect(db_path)
            conn.row_factory = sqlite3.Row
            cursor = conn.cursor()
            cursor.execute(query)

            results = []
            for row in cursor.fetchall():
                contract_data = dict(row)
                contract_data["source"] = f"sqlite:{db_path}"
                contract_data["extraction_date"] = datetime.now().isoformat()
                contract_data.update(self._standardize_contract_data(contract_data))
                results.append(contract_data)

            return results

        finally:
            if conn:
                conn.close()

    def _standardize_contract_data(self, raw_data: Dict[str, Any]) -> Dict[str, Any]:
        """
        Standardize contract data fields

        Args:
            raw_data: Raw extracted data

        Returns:
            Standardized contract data
        """
        standardized = {}

        # Map common field variations to standard names
        field_mappings = {
            "contract_id": ["contract_id", "id", "contract_number", "contract_no"],
            "contract_date": ["contract_date", "date", "created_date", "start_date"],
            "contract_amount": ["amount", "value", "contract_value", "total_amount"],
            "client_name": ["client", "customer", "client_name", "customer_name"],
            "vendor_name": ["vendor", "supplier", "contractor", "vendor_name"],
            "status": ["status", "contract_status", "state"],
            "description": ["description", "details", "contract_description"],
        }

        for standard_field, possible_fields in field_mappings.items():
            for field in possible_fields:
                if field in raw_data and raw_data[field] is not None:
                    value = raw_data[field]

                    # Type conversion and validation
                    if standard_field == "contract_amount" and isinstance(
                        value, (str, int, float)
                    ):
                        try:
                            if isinstance(value, str):
                                value = value.replace(",", "").replace("$", "").strip()
                            standardized[standard_field] = float(value)
                        except:
                            standardized[f"{standard_field}_raw"] = value
                    elif standard_field == "contract_date" and isinstance(value, str):
                        try:
                            parsed_date = date_parser.parse(value)
                            standardized[standard_field] = parsed_date.isoformat()
                        except:
                            standardized[f"{standard_field}_raw"] = value
                    else:
                        standardized[standard_field] = value
                    break

        return standardized


def extract_contracts(
    sources: Union[str, List[str], Dict[str, Any]],
    source_type: str = "auto",
    config: Optional[Dict] = None,
) -> List[Dict[str, Any]]:
    """
    Main function to extract contracts from various sources

    Args:
        sources: Single source path/query or list of sources, or configuration dict
        source_type: Type of source ('pdf', 'docx', 'txt', 'csv', 'database', 'auto')
        config: Configuration dictionary

    Returns:
        List of dictionaries containing contract information
    """
    extractor = ContractExtractor(config)
    results = []

    # Handle single source
    if isinstance(sources, str):
        sources = [sources]

    # Handle configuration dictionary
    if isinstance(sources, dict):
        return _extract_from_config(sources, extractor)

    for source in sources:
        try:
            if source_type == "auto":
                source_type = _detect_source_type(source)

            if source_type == "pdf":
                result = extractor.extract_from_pdf(source)
                results.append(result)
            elif source_type == "docx":
                result = extractor.extract_from_docx(source)
                results.append(result)
            elif source_type == "txt":
                result = extractor.extract_from_text_file(source)
                results.append(result)
            elif source_type == "csv":
                csv_results = extractor.extract_from_csv(source)
                results.extend(csv_results)
            else:
                logger.warning(
                    f"Unsupported source type: {source_type} for source: {source}"
                )

        except Exception as e:
            logger.error(f"Error processing source {source}: {str(e)}")
            results.append(
                {
                    "source": source,
                    "error": str(e),
                    "extraction_date": datetime.now().isoformat(),
                }
            )

    return results


def _detect_source_type(source: str) -> str:
    """Detect source type based on file extension"""
    if not isinstance(source, str):
        return "unknown"

    extension = Path(source).suffix.lower()

    if extension == ".pdf":
        return "pdf"
    elif extension in [".docx", ".doc"]:
        return "docx"
    elif extension == ".txt":
        return "txt"
    elif extension == ".csv":
        return "csv"
    else:
        return "txt"  # Default to text


def _extract_from_config(
    config: Dict[str, Any], extractor: ContractExtractor
) -> List[Dict[str, Any]]:
    """Extract contracts based on configuration dictionary"""
    results = []

    # Handle file sources
    if "files" in config:
        for file_config in config["files"]:
            file_path = file_config["path"]
            file_type = file_config.get("type", "auto")

            if file_type == "auto":
                file_type = _detect_source_type(file_path)

            try:
                if file_type == "pdf":
                    result = extractor.extract_from_pdf(file_path)
                elif file_type == "docx":
                    result = extractor.extract_from_docx(file_path)
                elif file_type == "txt":
                    result = extractor.extract_from_text_file(file_path)
                elif file_type == "csv":
                    csv_results = extractor.extract_from_csv(
                        file_path, file_config.get("column_mapping")
                    )
                    results.extend(csv_results)
                    continue
                else:
                    logger.warning(f"Unsupported file type: {file_type}")
                    continue

                results.append(result)

            except Exception as e:
                logger.error(f"Error processing file {file_path}: {str(e)}")
                results.append(
                    {
                        "source": file_path,
                        "error": str(e),
                        "extraction_date": datetime.now().isoformat(),
                    }
                )

    # Handle database sources
    if "databases" in config:
        for db_config in config["databases"]:
            try:
                db_results = extractor.extract_from_database(
                    db_config["query"],
                    db_config["connection_params"],
                    db_config.get("type", "postgresql"),
                )
                results.extend(db_results)

            except Exception as e:
                logger.error(f"Error processing database query: {str(e)}")
                results.append(
                    {
                        "source": f"database_query",
                        "error": str(e),
                        "extraction_date": datetime.now().isoformat(),
                    }
                )

    return results


def validate_contract_data(contract_data: Dict[str, Any]) -> Dict[str, Any]:
    """
    Validate and clean contract data

    Args:
        contract_data: Contract data dictionary

    Returns:
        Validated contract data with validation results
    """
    validation_results = {"is_valid": True, "errors": [], "warnings": []}

    # Check for required fields
    required_fields = ["source"]
    for field in required_fields:
        if field not in contract_data or not contract_data[field]:
            validation_results["errors"].append(f"Missing required field: {field}")
            validation_results["is_valid"] = False

    # Validate contract amount
    if "contract_amount" in contract_data:
        try:
            amount = float(contract_data["contract_amount"])
            if amount < 0:
                validation_results["warnings"].append("Contract amount is negative")
        except (ValueError, TypeError):
            validation_results["errors"].append("Invalid contract amount format")

    # Validate dates
    if "contract_date" in contract_data:
        try:
            date_parser.parse(contract_data["contract_date"])
        except (ValueError, TypeError):
            validation_results["errors"].append("Invalid contract date format")

    # Validate email format
    if "emails" in contract_data:
        email_pattern = re.compile(
            r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b"
        )
        for email in contract_data["emails"]:
            if not email_pattern.match(email):
                validation_results["warnings"].append(f"Invalid email format: {email}")

    contract_data["validation"] = validation_results
    return contract_data


if __name__ == "__main__":
    # Example usage
    sample_config = {
        "files": [
            {"path": "sample_contract.pdf", "type": "pdf"},
            {"path": "contracts.csv", "type": "csv"},
        ],
        "databases": [
            {
                "type": "postgresql",
                "connection_params": {
                    "host": "localhost",
                    "database": "contracts",
                    "user": "user",
                    "password": "password",
                },
                "query": "SELECT * FROM contracts WHERE status = %s",
                "params": ["active"],
            }
        ],
    }

    try:
        contracts = extract_contracts(sample_config)
        print(f"Extracted {len(contracts)} contracts")
        for contract in contracts:
            validated_contract = validate_contract_data(contract)
            print(f"Contract: {validated_contract.get('contract_number', 'Unknown')}")
    except Exception as e:
        print(f"Error: {str(e)}")
